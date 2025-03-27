from docx import Document

# Create a new Word document
doc = Document()

# Add content to the document
doc.add_heading('Robot Agent Pathfinding and Delivery Simulation', level=1)

doc.add_paragraph(
    "The `robot_agent.py` script simulates a robot navigating a warehouse to deliver packages while avoiding obstacles. "
    "Below is a concise explanation of the process:"
)

doc.add_heading('1. Warehouse Setup', level=2)
doc.add_paragraph(
    "- Grid Representation: The warehouse is an N x M grid.\n"
    "- Elements: Obstacles ('X'), Packages ('P'), Drop-off Points ('D'), and Robot ('R') are randomly placed.\n"
    "- Initialization: The `initialize_warehouse` function ensures no overlap between elements."
)

doc.add_heading('2. Pathfinding (BFS)', level=2)
doc.add_paragraph(
    "- Algorithm: Breadth-First Search (BFS) is used to find the shortest path between the robot's position and its target.\n"
    "- Mechanism: BFS explores all valid moves while avoiding obstacles and revisited nodes.\n"
    "- Penalty: Crossing obstacles adds a penalty to the score."
)

doc.add_heading('3. Package Delivery Process', level=2)
doc.add_paragraph(
    "1. The robot moves to the package location using BFS.\n"
    "2. It then moves to the corresponding drop-off point.\n"
    "3. This process repeats for all packages."
)

doc.add_heading('4. Scoring System', level=2)
doc.add_paragraph(
    "- Cost: Steps taken by the robot.\n"
    "- Reward: Each successful delivery adds 10 points.\n"
    "- Penalty: Each obstacle crossed reduces 5 points.\n"
    "- Final Score: Calculated as Total Reward - Total Cost."
)

doc.add_heading('5. Execution Example', level=2)
doc.add_paragraph(
    "The script generates a random warehouse layout, simulates package deliveries, and outputs:\n"
    "Final Score: 15\n"
    "Total Cost: 10\n"
    "Total Reward: 25\n"
    "Penalty Cost: 0"
)

# Save the document
file_name = 'Robot_Agent_Report.docx'
doc.save(file_name)

print(f"Report saved as {file_name}")